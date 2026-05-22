from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Add title
title = doc.add_heading('Roots & Routes - Project Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add intro
doc.add_heading('Project Overview', level=1)
doc.add_paragraph(
    'Roots & Routes is a personality + career-fit assessment web application built with Next.js 15.5 and React 19. '
    'It serves as an in-funnel replacement for PRISM in the Secure Steps flow, providing students with a comprehensive '
    '~20-page personalized report based on ~45 scenario-based assessment questions.'
)

doc.add_paragraph('Live deployment: https://roots-and-routes-rho.vercel.app')

# Stack
doc.add_heading('Technology Stack', level=1)
tech_items = [
    ('Frontend Framework', 'Next.js 15.5 (App Router) with React 19'),
    ('Language', 'TypeScript (strict mode)'),
    ('Styling', 'Tailwind CSS 3 + custom utility classes'),
    ('State Management', 'Zustand with localStorage persistence'),
    ('Animation', 'framer-motion for card transitions and radar animations'),
    ('PDF Generation', '@react-pdf/renderer (client-side rendering)'),
    ('AI Integration', '@anthropic-ai/sdk for adaptive follow-ups and rewrites'),
    ('UI Components', 'Radix UI (dialog, popover, progress, slot)'),
    ('Icons', 'lucide-react'),
    ('Runtime', 'Node 20+'),
]

for tech, desc in tech_items:
    doc.add_paragraph(f'{tech}: {desc}', style='List Bullet')

# Assessment Flow
doc.add_heading('Assessment Flow & Sections', level=1)
doc.add_paragraph(
    'The assessment is divided into three main sections that guide students through a comprehensive evaluation:'
)

sections = {
    'Context (10 questions)': 'Students provide background information including budget, geography, family situation, ambition level, timeline, and 5-year aspirations.',
    'Roots (20 questions)': 'Universal personality assessment covering 6 core dimensions using scenario-based questions.',
    'Teaser': 'Halfway reveal screen showing the identified archetype.',
    'Routes (8 + 1 engagement check)': 'Discipline-specific scenarios tailored to the selected field of interest, plus an engagement check.',
}

for section, description in sections.items():
    doc.add_paragraph(f'{section}: {description}', style='List Bullet')

# Six Dimensions
doc.add_heading('Six Core Dimensions', level=2)
dimensions = [
    ('Decision Style', 'Intuitive ↔ Deliberate'),
    ('Energy', 'Introvert ↔ Extrovert'),
    ('Structure', 'Open-ended ↔ Structured'),
    ('Risk Tolerance', 'Cautious ↔ Bold'),
    ('Social Mode', 'Independent ↔ Collaborative'),
    ('Drive', 'Reactive ↔ Proactive'),
]

for dim, spectrum in dimensions:
    doc.add_paragraph(f'{dim}: {spectrum}', style='List Bullet')

# Archetypes
doc.add_heading('Six Archetypes', level=2)
archetypes = [
    'The Builder: "You\'d rather ship at 80% than discuss at 100%." - Output-driven, first-mover mentality.',
    'The Strategist: "You see two moves ahead before most people see one." - Deliberate decision-maker, trusts structure.',
    'The Connector: "You\'re the hub in a network." - Builds relationships, cross-disciplinary collaborator.',
    'The Maverick: "You follow the signal, not the map." - Independent thinker, frontier explorer.',
    'The Anchor: "You\'re the stability others build on." - Reliability-focused, distributed systems thinker.',
    'The Explorer: "Not all those who wander are lost." - Polymath, adaptable, boundary-crosser.',
]

for archetype in archetypes:
    doc.add_paragraph(archetype, style='List Bullet')

# Report Contents
doc.add_heading('Report Output (~20 pages)', level=1)
report_items = [
    'Six dimensional readings (rendered as radar chart)',
    '30+ derived stats (strengths, growth edges, work environment preference)',
    'Eight work aptitudes (rendered as radar chart)',
    '18 environment-fit predictions',
    '26 career-development traits across 6 categories',
    'Top course recommendations from 50+ course catalogue, ranked by archetype fit',
    'Adjacent / alternative course pathways',
    'Drivers: 5 personalized narrative insights',
    'Future Day: 3-paragraph narrative of typical day in matched field',
    'Parent Letter: Short formal letter for family',
    'Next Steps: Actionable recommendations',
    'Thank-you closer',
]

for item in report_items:
    doc.add_paragraph(item, style='List Bullet')

# Optional AI Layers
doc.add_heading('Optional AI Layers (with Anthropic API Key)', level=2)
ai_features = [
    'Adaptive Follow-up Questions: Claude generates follow-ups tailored to student answers',
    'Reaction-Driven Rewrites: Students can react to questions (positive/negative) for LLM rephrasing',
]

for feature in ai_features:
    doc.add_paragraph(feature, style='List Bullet')

# Project Structure
doc.add_heading('Project Structure & Key Files', level=1)

doc.add_heading('App Directory (Pages)', level=2)
app_files = [
    ('layout.tsx', 'Root layout with Inter + JetBrains Mono fonts'),
    ('globals.css', 'Tailwind layers + custom utility classes (.panel, .mono-eyebrow)'),
    ('page.tsx', 'Landing page - hero + sign-up form with 13 disciplines + course picker'),
    ('assessment/page.tsx', 'Main assessment orchestrator (state, navigation, reactions, probes)'),
    ('teaser/page.tsx', 'Halfway reveal screen showing identified archetype'),
    ('report/page.tsx', 'Report page - header, download PDF, end session'),
    ('thank-you/page.tsx', 'End-of-session screen - clears store, auto-redirects'),
    ('api/probe/route.ts', 'Server route: Claude probe (deeper/rephrase modes) or skip'),
]

for file, desc in app_files:
    doc.add_paragraph(f'{file}: {desc}', style='List Bullet')

doc.add_heading('Components Directory', level=2)
component_files = [
    ('report-view.tsx', 'HTML report rendering - cover, drivers, dimensions, stats, recs'),
    ('report-visuals.tsx', 'RadarChart & EngagementDial SVG components'),
    ('node-graph/assessment-graph.tsx', 'History rail + viewport-centered stage + transitions'),
    ('node-graph/question-card.tsx', 'Question card UI - single/multi/short-text + reactions'),
    ('node-graph/trunk-chip.tsx', 'Collapsed previous-answer chip (clickable to edit)'),
    ('node-graph/reaction-bar.tsx', '5-emoji reaction bar (Apple emoji images via CDN)'),
    ('ui/button.tsx', 'Button variants: outline/solid/ghost/electric'),
    ('ui/logo.tsx', 'Secure Steps pinwheel mark (SVG + img fallback'),
    ('ui/photo-upload.tsx', 'Circular drop-zone with canvas-based compression'),
]

for file, desc in component_files:
    doc.add_paragraph(f'{file}: {desc}', style='List Bullet')

doc.add_heading('Lib Directory (Core Logic)', level=2)
lib_files = [
    ('types.ts', 'Core type definitions (Question, Answer, Archetype, ReportData, StudentProfile)'),
    ('store.ts', 'Zustand store with localStorage persist - assessment state management'),
    ('flow.ts', 'planNext() - decides what question comes after current (anchor/probe/transition)'),
    ('image.ts', 'compressToAvatar() - canvas-based JPEG compression for photos'),
    ('anthropic.ts', 'Anthropic SDK client wrapper'),
    ('archetype.ts', 'computeDimensionScores() & pickArchetype() - scoring and archetype assignment'),
    ('course-catalog.ts', '50+ courses across 13 disciplines with careers + dimension weights'),
    ('report-data.ts', 'Big derivation engine - stats, work prefs, aptitudes, env fit, traits, drivers, letter'),
    ('report-pdf.tsx', '@react-pdf/renderer report - ~20 pages with radar + engagement dial'),
    ('utils.ts', 'cn() className helper'),
    ('question-bank/context.ts', '10 context questions (budget, geo, family, dream, timeline)'),
    ('question-bank/roots.ts', '20 Roots anchor scenarios'),
    ('question-bank/routes-bca.ts', 'Tech/CS Routes bank + engagement check'),
]

for file, desc in lib_files:
    doc.add_paragraph(f'{file}: {desc}', style='List Bullet')

# Key Functions & Methods
doc.add_heading('Critical Functions & Methods', level=1)

doc.add_heading('State Management (lib/store.ts)', level=2)
store_methods = [
    ('setProfile()', 'Sets the StudentProfile (name, email, discipline, optional course & photo)'),
    ('answer()', 'Submit or update an answer for a question'),
    ('react()', 'Record a student\'s reaction (fire/love/think/meh/bored) to a question'),
    ('setTrunk()', 'Replace the in-progress trunk (ordered list of question IDs)'),
    ('pushAdaptive()', 'Append an adaptive question generated by /api/probe'),
    ('setRewrite()', 'Store a rewrite for a question (negative-reaction rephrase)'),
    ('setSection()', 'Move to a section (context/roots/teaser/routes/report)'),
    ('computeArchetype()', 'Calculate archetype from current Roots answers'),
    ('reset()', 'Clear everything for "start over"'),
    ('getQuestion()', 'Retrieve question by ID with rewrite applied'),
    ('allQuestions()', 'Get all questions known so far (banks + adaptives + rewrites)'),
]

for method, desc in store_methods:
    doc.add_paragraph(f'{method}: {desc}', style='List Bullet')

doc.add_heading('Scoring & Archetypes (lib/archetype.ts)', level=2)
archetype_methods = [
    ('computeDimensionScores()', 'Sum option scores, normalize each dimension to -100..+100 scale'),
    ('pickArchetype()', 'Match student dimension scores to best-fit archetype of 6 options'),
]

for method, desc in archetype_methods:
    doc.add_paragraph(f'{method}: {desc}', style='List Bullet')

doc.add_heading('Report Generation (lib/report-data.ts)', level=2)
doc.add_paragraph(
    'buildReport() is the massive derivation engine that computes all report sections from profile, archetype, and answers:'
)

report_methods = [
    'buildStats() - 30+ derived statistics',
    'buildWorkPreferences() - 15 work environment preferences',
    'buildWorkAptitudes() - 8 aptitudes with 0-100 scores',
    'buildEnvironmentFit() - 18 environment compatibility scores',
    'buildCareerTraits() - 26 traits across 6 categories (Communication, Problem-solving, etc.)',
    'buildDrivers() - 5 personalized narrative drivers',
    'buildFutureDay() - 3-paragraph "typical day" narrative',
    'buildTopDegrees() - Top course recommendations sorted by fit',
    'buildCommonCareerPaths() - Career trajectory examples',
    'buildParentLetter() - Formal letter for family',
]

for method in report_methods:
    doc.add_paragraph(method, style='List Bullet')

doc.add_heading('Question Flow (lib/flow.ts)', level=2)
flow_methods = [
    ('planNext()', 'Decide what comes after current question: next anchor, adaptive probe, transition, or complete'),
    ('anchorsForSection()', 'Get all anchor questions for a given section'),
]

for method, desc in flow_methods:
    doc.add_paragraph(f'{method}: {desc}', style='List Bullet')

# Question Types & Data Structures
doc.add_heading('Question & Answer Types', level=1)

doc.add_heading('Question Interface', level=2)
question_props = [
    ('id', 'Unique identifier'),
    ('section', 'context | roots | teaser | routes | report'),
    ('kind', 'anchor | adaptive | context | engagement'),
    ('type', 'single_choice | multi_choice | short_text | single_select_chips'),
    ('category', 'Tag shown above prompt (e.g., "DECISION STYLE")'),
    ('prompt', 'Main question text'),
    ('hint', 'Optional helper text'),
    ('options', 'Array of Option objects with scores & tags'),
    ('parentId', 'For adaptive probes - branched from question ID'),
    ('dimension', 'Primary dimension (decision_style, energy, structure, risk, social, drive)'),
]

for prop, desc in question_props:
    doc.add_paragraph(f'{prop}: {desc}', style='List Bullet')

doc.add_heading('Answer Interface', level=2)
answer_props = [
    ('questionId', 'Reference to the Question'),
    ('optionIds', 'Array of selected option IDs (for multi-choice)'),
    ('text', 'Free-text answer for short_text questions'),
    ('reaction', 'Student reaction: fire | love | think | meh | bored'),
    ('answeredAt', 'Timestamp in milliseconds'),
]

for prop, desc in answer_props:
    doc.add_paragraph(f'{prop}: {desc}', style='List Bullet')

doc.add_heading('Reaction System', level=2)
reactions = [
    ('🔥 fire', 'Positive reaction - probe deeper'),
    ('❤️ love', 'Positive reaction - felt seen, probe deeper'),
    ('🤔 think', 'Neutral reaction - made me think'),
    ('😐 meh', 'Negative reaction - rephrase this'),
    ('😴 bored', 'Negative reaction - rephrase this'),
]

for reaction, tone in reactions:
    doc.add_paragraph(f'{reaction}: {tone}', style='List Bullet')

# Data Flow & Architecture
doc.add_heading('Data Flow & Architecture', level=1)

doc.add_heading('State Persistence', level=2)
doc.add_paragraph(
    'Assessment state lives in lib/store.ts (Zustand) and is persisted to localStorage under key "roots-and-routes-assessment". '
    'The schema is defined by AssessmentState interface in lib/types.ts. Currently no backend database; designed for easy migration to Supabase (~30 lines in store.ts).'
)

doc.add_heading('Assessment Flow', level=2)
doc.add_paragraph(
    '1. Student completes sign-up form (name, email, discipline, optional course)'
)
doc.add_paragraph(
    '2. Context section: 10 questions establishing background (budget, geo, family, ambition, timeline, 5-year dream)'
)
doc.add_paragraph(
    '3. Roots section: 20 personality anchor questions with 3 interleaved adaptive probes (if API key configured)'
)
doc.add_paragraph(
    '4. Teaser: Archetype reveal showing matched personality type'
)
doc.add_paragraph(
    '5. Routes section: 8 discipline-specific scenarios + 1 engagement check, with 2 interleaved adaptive probes'
)
doc.add_paragraph(
    '6. Report: Full ~20-page personalized report with stats, recommendations, narratives'
)
doc.add_paragraph(
    '7. Thank you: Session clearance and redirect to home'
)

doc.add_heading('Archetype Scoring Algorithm', level=2)
doc.add_paragraph(
    'Each question option has scores across 6 dimensions (range -2..+2). As students answer:'
)
doc.add_paragraph(
    '1. Raw scores are summed for each dimension',
    style='List Number'
)
doc.add_paragraph(
    '2. Scores are normalized to -100..+100 scale',
    style='List Number'
)
doc.add_paragraph(
    '3. All 6 archetypes are evaluated using match function (weighted dimension combinations)',
    style='List Number'
)
doc.add_paragraph(
    '4. Highest-scoring archetype is selected',
    style='List Number'
)

doc.add_heading('Adaptive Question Generation (/api/probe)', level=2)
doc.add_paragraph(
    'When a probe slot is triggered (at specific anchor indices), the frontend calls /api/probe with:'
)
doc.add_paragraph('- Parent question and student answer', style='List Bullet')
doc.add_paragraph('- Question trunk history', style='List Bullet')
doc.add_paragraph('- All student answers so far', style='List Bullet')
doc.add_paragraph('- Section context (roots or routes)', style='List Bullet')

doc.add_paragraph(
    'Claude generates a follow-up question (deeper or rephrase mode) tailored to the student\'s response. '
    'Negative reactions trigger rephrase mode; positive reactions trigger deeper mode.'
)

# Important Keywords & Concepts
doc.add_heading('Important Keywords & Concepts', level=1)

keywords = {
    'Trunk': 'Ordered list of question IDs currently active in the assessment flow',
    'Anchor': 'Base questions from the question bank (not adaptive)',
    'Adaptive': 'AI-generated probe questions branching off anchor answers',
    'Discipline': 'Field of interest (13 categories: tech_cs, business, creative, etc.)',
    'Reaction': 'Student emotional response to a question (fire/love/think/meh/bored)',
    'Rewrite': 'Alternative phrasing of a question triggered by negative reaction',
    'Engagement Readout': 'Score measuring how committed student feels to their selected field',
    'Dimension': 'One of 6 personality axes (decision_style, energy, structure, risk, social, drive)',
    'Archetype': 'One of 6 personality types (Builder, Strategist, Connector, Maverick, Anchor, Explorer)',
    'Niche': 'Role category within a discipline (e.g., "DevOps Engineer", "Product Manager")',
    'Route Cluster': 'Grouped roles within discipline sorted by fit to student profile',
    'Work Preference': 'Derived trait about ideal work environment (remote, team size, impact, etc.)',
    'Work Aptitude': 'Capability score (0-100) in areas like "Leadership", "Creative Problem-Solving"',
    'Environment Fit': 'Compatibility score (0-100) with different work environments',
    'Career Trait': 'Personality characteristic in 6 categories (Communication, Resilience, etc.)',
}

for term, definition in keywords.items():
    doc.add_paragraph(f'{term}: {definition}', style='List Bullet')

# Configuration & Customization
doc.add_heading('Configuration & Customization', level=1)

doc.add_heading('Editing the Question Bank', level=2)
doc.add_paragraph(
    'Questions are split by section and stored as TypeScript arrays:'
)
doc.add_paragraph('- lib/question-bank/context.ts: 10 context questions', style='List Bullet')
doc.add_paragraph('- lib/question-bank/roots.ts: 20 Roots anchor scenarios', style='List Bullet')
doc.add_paragraph('- lib/question-bank/routes-bca.ts: Routes bank for tech/CS discipline', style='List Bullet')

doc.add_paragraph(
    'Each question has options with dimension scores. Modify scores to rebalance how questions contribute to archetype matching.'
)

doc.add_heading('Course Catalogue', level=2)
doc.add_paragraph(
    'lib/course-catalog.ts contains 50+ courses across 13 disciplines with metadata:'
)
doc.add_paragraph('- Course name, description, careers', style='List Bullet')
doc.add_paragraph('- 6 dimension weights for fit calculation', style='List Bullet')

doc.add_heading('Environment Variables', level=2)
doc.add_paragraph(
    'Set ANTHROPIC_API_KEY to enable adaptive questions and reaction-driven rewrites. '
    'Without it, the app still works but only uses static questions.'
)

# Deployment & Operations
doc.add_heading('Build & Deployment', level=1)

doc.add_heading('Development', level=2)
doc.add_paragraph('npm run dev - Start Next.js dev server on http://localhost:3000', style='List Bullet')

doc.add_heading('Production Build', level=2)
doc.add_paragraph('npm run build - Compile and optimize', style='List Bullet')
doc.add_paragraph('npm run start - Run production server', style='List Bullet')

doc.add_heading('Linting & Type Checking', level=2)
doc.add_paragraph('npm run lint - Run ESLint', style='List Bullet')
doc.add_paragraph('npm run typecheck - Run TypeScript compiler (tsc --noEmit)', style='List Bullet')

doc.add_heading('Live Deployment', level=2)
doc.add_paragraph('Deployed on Vercel: https://roots-and-routes-rho.vercel.app')

# Important Notes
doc.add_heading('Important Implementation Notes', level=1)

notes = [
    'localStorage is used for state persistence - no backend DB yet. Easy to swap to Supabase when needed.',
    'PDF generation happens client-side using @react-pdf/renderer - no server function needed.',
    'Adaptive question probes are optional (require ANTHROPIC_API_KEY); the app gracefully skips them if not configured.',
    'All dimensions are normalized to -100..+100 scale for consistent reporting.',
    'The trunk (question order) is computed dynamically based on current section and archetype state.',
    'Reactions are stored per-question and can trigger LLM rewrites for negative responses.',
    'The archetype reveal (Teaser page) happens after Roots section completion.',
    'Routes questions are currently discipline-generic; per-discipline routes banks planned for Phase 2.',
    'Parent Letter is a formal, context-aware message generated for family members.',
    'Report contains 120+ distinct measurements across all categories.',
    'Framer Motion handles smooth card transitions and radar polygon animations.',
]

for note in notes:
    doc.add_paragraph(note, style='List Bullet')

# Save document
doc.save('Secure_Steps_Project_Documentation.docx')
print('✅ Document created: Secure_Steps_Project_Documentation.docx')
